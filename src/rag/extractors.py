"""extractors.py -- text extraction for PDF / DOCX / DOC.

Vendored from the toolkit's doc_reader.py + pdf_evidence_locator.py so docrag
has no external-package dependency beyond PyPDF2 + python-docx. Plain-text
formats (.txt/.md) are read directly by extract.py.

Public API:
    sanitize_ascii(text) -> str
    extract_pdf_pages(path) -> list[(page_no, text)]   # 1-indexed
    extract_docx(path) -> str
    extract_doc(path) -> str        # legacy binary .doc, Latin-1 fallback
    extract_html(path) -> dict      # {text, book, chapter} -- tags stripped
"""

from __future__ import annotations

import html as _html
import os
import re
import sys


_REPLACEMENTS = {
    "‘": "'",   # left single curly quote
    "’": "'",   # right single curly quote
    "“": '"',   # left double curly quote
    "”": '"',   # right double curly quote
    "—": "--",  # em dash
    "–": "-",   # en dash
    "…": "...",  # ellipsis
    "•": "-",   # bullet
    " ": " ",   # non-breaking space
}


def sanitize_ascii(text: str) -> str:
    """Replace non-ASCII characters with safe equivalents.

    Single chokepoint for ASCII enforcement -- every extractor pipes output
    through this so downstream chunking / embedding sees consistent bytes.
    """
    if not text:
        return ""
    for old, new in _REPLACEMENTS.items():
        text = text.replace(old, new)
    # Strip control characters (keep \t \n \r).
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    out = []
    for ch in text:
        out.append(ch if ord(ch) <= 0x7E else "?")
    return "".join(out)


# -- PDF -----------------------------------------------------------------------


def extract_pdf_pages(pdf_path: str) -> list[tuple[int, str]]:
    """Extract text from every page, returning ``(page_no, text)`` tuples.

    page_no is 1-indexed. Pages with no extractable text are included as
    empty strings so page indexing stays correct. Output is ASCII-sanitized.
    Returns ``[]`` (with a stderr warning) when the PDF can't be opened.
    """
    try:
        import PyPDF2  # type: ignore
    except ImportError:
        sys.stderr.write("[extract] PyPDF2 not installed\n")
        return []

    try:
        reader = PyPDF2.PdfReader(pdf_path)
    except Exception as exc:  # noqa: BLE001
        sys.stderr.write("[extract] cannot open PDF: %s\n" % exc)
        return []

    if reader.is_encrypted:
        # Many PDFs are "encrypted" with an empty owner password; try that.
        try:
            reader.decrypt("")
        except Exception:  # noqa: BLE001
            sys.stderr.write("[extract] PDF is encrypted: %s\n" % pdf_path)
            return []

    try:
        n_pages = len(reader.pages)
    except Exception as exc:  # noqa: BLE001 -- e.g. encrypted/AES parse failure
        sys.stderr.write("[extract] cannot enumerate PDF pages (%s): %s\n"
                         % (exc, pdf_path))
        return []
    pages: list[tuple[int, str]] = []
    for i in range(n_pages):
        try:
            text = reader.pages[i].extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        pages.append((i + 1, sanitize_ascii(text)))
    return pages


# -- PDF (structure-aware, layout model) ---------------------------------------

# Cosmetic only: split an ALREADY-identified heading into (number, title) so
# citations can read "§ 160D-1110 Building permits". This does NOT find
# headings -- the layout model does that. number=None when no leading
# code/statute-style token is present (the whole heading becomes the title).
_HEADING_SPLIT_RE = re.compile(
    r"^\s*(?:§\s*)?([A-Z]{0,4}\d[\w.\-]*?)\.?\s+(.*\S)?\s*$")


def _split_heading(text: str):
    text = " ".join((text or "").split())
    m = _HEADING_SPLIT_RE.match(text)
    if m and any(ch.isdigit() for ch in m.group(1)):
        return m.group(1), (m.group(2) or None)
    return None, (text or None)


def extract_pdf_docling(pdf_path: str):
    """Recover a PDF's heading hierarchy with a trained layout model (Docling).

    Returns a ``sections`` list in the SAME shape as ``extract_html`` --
    ``[{section_id, section_number, section_title, parent_section_id, level,
    own_text, full_text, has_table, page}]`` -- so the caller routes it through
    the existing structure-aware (HTML) chunking path. Returns ``None`` when
    Docling is unavailable, conversion fails, or no headings are detected (the
    caller then falls back to page-chunking). Text is raw; caller sanitizes.
    """
    # On Windows the HF cache uses symlinks by default, which fails without
    # admin/developer mode (WinError 1314). Force plain-copy caching.
    os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS", "1")
    os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")

    try:
        from docling.document_converter import DocumentConverter  # type: ignore
    except Exception as exc:  # noqa: BLE001 -- dep absent
        sys.stderr.write("[extract] docling unavailable (%s); page fallback\n" % exc)
        return None

    # Disable OCR: these are born-digital code/statute PDFs with a real text
    # layer, so OCR is wasteful (slow + extra model downloads) and less
    # accurate than the embedded text. Keep table-structure recovery on.
    try:
        from docling.datamodel.base_models import InputFormat  # type: ignore
        from docling.datamodel.pipeline_options import PdfPipelineOptions  # type: ignore
        from docling.document_converter import PdfFormatOption  # type: ignore
        _opts = PdfPipelineOptions(do_ocr=False, do_table_structure=True)
        _converter = DocumentConverter(
            format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=_opts)})
    except Exception as exc:  # noqa: BLE001 -- API drift -> defaults
        sys.stderr.write("[extract] docling options unavailable (%s); defaults\n" % exc)
        _converter = DocumentConverter()

    def _label(item) -> str:
        lab = getattr(item, "label", "")
        return getattr(lab, "value", str(lab)).lower()

    def _page_of(item):
        prov = getattr(item, "prov", None) or []
        if prov:
            return getattr(prov[0], "page_no", None)
        return None

    def _table_md(item, text):
        try:
            return item.export_to_dataframe().to_markdown(index=False)
        except Exception:  # noqa: BLE001 -- older docling / no pandas
            return text or ""

    # The layout model renders each page to an image; doing all pages at once
    # exhausts memory on large statutes (std::bad_alloc). Process in bounded
    # page batches to cap peak memory. A section that spans a batch boundary
    # keeps accumulating because the heading stack persists across batches, and
    # page numbers from page_range are absolute. Normalize items to plain tuples
    # (label, text, page, level) so docling objects are freed between batches.
    try:
        import PyPDF2  # type: ignore
        n_pages = len(PyPDF2.PdfReader(pdf_path).pages)
    except Exception:  # noqa: BLE001
        n_pages = None

    BATCH = 16
    items: list[tuple] = []

    def _collect(doc):
        for entry in doc.iterate_items():
            it = entry[0] if isinstance(entry, tuple) else entry
            lab = _label(it)
            text = _table_md(it, getattr(it, "text", None)) if lab == "table" \
                else getattr(it, "text", None)
            items.append((lab, text, _page_of(it), getattr(it, "level", None)))

    try:
        if n_pages and n_pages > BATCH:
            for start in range(1, n_pages + 1, BATCH):
                end = min(start + BATCH - 1, n_pages)
                try:
                    d = _converter.convert(pdf_path, page_range=(start, end)).document
                    _collect(d)
                    del d
                except Exception as exc:  # noqa: BLE001 -- skip a bad batch, keep going
                    sys.stderr.write("[extract] docling batch %d-%d failed (%s)\n"
                                     % (start, end, exc))
        else:
            _collect(_converter.convert(pdf_path).document)
    except Exception as exc:  # noqa: BLE001
        sys.stderr.write("[extract] docling convert failed (%s); page fallback\n" % exc)
        return None

    if not items:
        return None

    # Walk items in reading order; headings open sections, text/tables append to
    # the current section. A stack keyed by heading level gives parent links.
    nodes: dict[str, dict] = {}
    order: list[str] = []
    stack: list[tuple[int, str]] = []  # (heading_level, section_id)
    counter = 0

    for lab, text, page, lvl in items:
        if lab in ("title", "section_header", "page_header"):
            level = lvl if isinstance(lvl, int) else (1 if lab != "title" else 0)
            counter += 1
            sid = "sec%04d" % counter
            num, title = _split_heading(text or "")
            while stack and stack[-1][0] >= level:
                stack.pop()
            parent_id = stack[-1][1] if stack else None
            nodes[sid] = {
                "section_id": sid, "section_number": num,
                "section_title": title, "parent_section_id": parent_id,
                "level": level, "own_text_parts": [], "has_table": False,
                "page": page,
            }
            order.append(sid)
            stack.append((level, sid))
        else:
            if not stack:
                continue  # preamble before the first heading -- skip
            cur = nodes[stack[-1][1]]
            if lab == "table":
                cur["has_table"] = True
            if text:
                cur["own_text_parts"].append(text)

    if not order:
        return None

    for sid in order:
        nodes[sid]["own_text"] = "\n".join(nodes[sid].pop("own_text_parts")).strip()

    # full_text = own + all descendants (depth-first, document order).
    children: dict[str, list[str]] = {}
    for sid in order:
        p = nodes[sid]["parent_section_id"]
        if p in nodes:
            children.setdefault(p, []).append(sid)

    def _full(sid: str) -> str:
        parts = [nodes[sid]["own_text"]] if nodes[sid]["own_text"] else []
        for ch in children.get(sid, []):
            ft = _full(ch)
            if ft:
                parts.append(ft)
        return "\n\n".join(parts)

    sections = []
    for sid in order:
        n = dict(nodes[sid])
        n["full_text"] = _full(sid)
        sections.append(n)
    return sections


# -- DOCX ----------------------------------------------------------------------


def extract_docx(filepath: str) -> str:
    """Extract text from a .docx file with light markdown structure.

    Returns raw (un-sanitized) text -- caller sanitizes. Raises RuntimeError
    if python-docx is missing or the file can't be opened, so the indexer can
    record the failure instead of crashing the process.
    """
    try:
        import docx  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx not installed") from exc

    try:
        document = docx.Document(filepath)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("cannot open DOCX: %s" % exc) from exc

    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            m = re.search(r"(\d+)", style_name)
            level = int(m.group(1)) if m else 1
            parts.append("%s %s" % ("#" * level, text))
        elif "List" in style_name:
            parts.append("- %s" % text)
        else:
            parts.append(text)

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip() or " "
                cells.append(cell_text.replace("\n", " "))
            rows.append(cells)
        if rows:
            parts.append("")
            parts.append("| %s |" % " | ".join(rows[0]))
            parts.append("| %s |" % " | ".join("---" for _ in rows[0]))
            for row_cells in rows[1:]:
                while len(row_cells) < len(rows[0]):
                    row_cells.append(" ")
                parts.append("| %s |" % " | ".join(row_cells[:len(rows[0])]))
            parts.append("")

    return "\n".join(parts)


# -- HTML ----------------------------------------------------------------------


_META_RE = re.compile(
    r'<meta\s+name="([^"]+)"\s+content="([^"]*)"', re.IGNORECASE)
_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"[ \t\f\v]+")
_BLANKS_RE = re.compile(r"\n\s*\n\s*\n+")
# Section number embedded in the ICC element id, e.g.
# "NCRC2024V1.0_Pt01_Ch01_SubCh01_SecR101.2.1" -> "R101.2.1".
_SEC_ID_NUM_RE = re.compile(r"_Sec([A-Za-z]?\d[\w.]*)")


def _meta(raw: str, name: str) -> str:
    for m in _META_RE.finditer(raw):
        if m.group(1).lower() == name.lower():
            return _html.unescape(m.group(2)).strip()
    return ""


def _clean_text(text: str) -> str:
    """Collapse intra-line whitespace, drop empty lines, cap blank runs."""
    text = _html.unescape(text or "")
    lines = [_WS_RE.sub(" ", ln).strip() for ln in text.splitlines()]
    text = "\n".join(ln for ln in lines if ln)
    return _BLANKS_RE.sub("\n\n", text)


def _table_to_markdown(table) -> str:
    """Render a <table> as GitHub-flavored markdown (header + rows)."""
    rows = []
    for tr in table.find_all("tr"):
        cells = tr.find_all(["th", "td"])
        if not cells:
            continue
        rows.append([" ".join(c.get_text(" ").split()) or " " for c in cells])
    if not rows:
        return ""
    ncol = max(len(r) for r in rows)
    rows = [r + [" "] * (ncol - len(r)) for r in rows]
    out = ["| " + " | ".join(rows[0]) + " |",
           "| " + " | ".join(["---"] * ncol) + " |"]
    for r in rows[1:]:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def _regex_strip(raw: str) -> str:
    return _clean_text(_html.unescape(_TAG_RE.sub(" ", raw)))


def extract_html(filepath: str) -> dict:
    """Parse an HTML file into a section tree + provenance.

    Built for the ICC-scraped code pages: each ``<section id="..._SecR105">``
    carries a stable id (breadcrumb), a ``.tg-number`` / ``.tg-title`` heading,
    nested child sections (the hierarchy), and ``<table>`` elements. Returns::

        {
          "book": str, "chapter": str,
          "sections": [ {section_id, section_number, section_title, level,
                         parent_section_id, own_text, full_text, has_table}, ...],
          "flat_text": str,   # fallback / whole-doc text
        }

    ``sections`` is empty when the document has no ``<section>`` markup -- the
    caller then falls back to ``flat_text`` + sliding-window chunking. Text is
    raw (un-sanitized); the caller pipes it through ``sanitize_ascii``.
    """
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
    except OSError as exc:
        raise RuntimeError("cannot read HTML: %s" % exc) from exc

    book = _meta(raw, "book")
    chapter = _meta(raw, "chapter")

    try:
        from bs4 import BeautifulSoup  # type: ignore
        try:
            soup = BeautifulSoup(raw, "lxml")
        except Exception:  # noqa: BLE001 -- lxml missing -> stdlib parser
            soup = BeautifulSoup(raw, "html.parser")
    except Exception as exc:  # noqa: BLE001 -- bs4 absent
        sys.stderr.write("[extract] bs4 unavailable (%s); regex strip\n" % exc)
        return {"book": book, "chapter": chapter, "sections": [],
                "flat_text": _regex_strip(raw)}

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    body = soup.body or soup

    sec_tags = body.find_all("section")
    if not sec_tags:
        # No structure -- whole-doc flat text (with tables as markdown inline).
        for tbl in body.find_all("table"):
            tbl.replace_with(soup.new_string("\n" + _table_to_markdown(tbl) + "\n"))
        return {"book": book, "chapter": chapter, "sections": [],
                "flat_text": _clean_text(body.get_text("\n"))}

    # Heading number/title for a section: the first .tg-number/.tg-title whose
    # nearest <section> ancestor is this section (i.e. not a child's heading).
    def _own_heading(sec):
        num = title = None
        for span in sec.select("span.tg-number, span.tg-title"):
            if span.find_parent("section") is not sec:
                continue
            txt = " ".join(span.get_text(" ").split())
            if "tg-number" in (span.get("class") or []) and num is None:
                num = txt
            elif "tg-title" in (span.get("class") or []) and title is None:
                title = txt
            if num is not None and title is not None:
                break
        if not num:
            m = _SEC_ID_NUM_RE.search(sec.get("id") or "")
            if m:
                num = m.group(1)
        return num, title

    # Process deepest sections first: capture own_text, render tables, then
    # detach the section so its parent's own_text excludes it.
    sec_tags.sort(key=lambda s: len(list(s.parents)), reverse=True)
    nodes: dict[str, dict] = {}
    order: list[str] = []
    for sec in sec_tags:
        sid = sec.get("id") or ""
        if not sid:
            continue
        parent_sec = sec.find_parent("section")
        parent_id = (parent_sec.get("id") or None) if parent_sec else None
        num, title = _own_heading(sec)
        for tbl in sec.find_all("table"):
            tbl.replace_with(soup.new_string("\n" + _table_to_markdown(tbl) + "\n"))
        has_table = "|" in (sec.get_text() or "") and "---" in (sec.get_text() or "")
        own = _clean_text(sec.get_text("\n"))
        nodes[sid] = {
            "section_id": sid, "section_number": num, "section_title": title,
            "parent_section_id": parent_id,
            "level": len(list(sec.parents)), "own_text": own,
            "has_table": has_table,
        }
        order.append(sid)
        sec.extract()  # remove from tree so ancestors' own_text excludes it

    # Restore document order (root-first) for stable downstream processing.
    order.reverse()

    # full_text = own_text + all descendant own_texts (depth-first, doc order).
    children: dict[str, list[str]] = {}
    for sid in order:
        p = nodes[sid]["parent_section_id"]
        if p in nodes:
            children.setdefault(p, []).append(sid)

    def _full_text(sid: str) -> str:
        parts = [nodes[sid]["own_text"]] if nodes[sid]["own_text"] else []
        for child in children.get(sid, []):
            ft = _full_text(child)
            if ft:
                parts.append(ft)
        return "\n\n".join(parts)

    sections = []
    for sid in order:
        n = dict(nodes[sid])
        n["full_text"] = _full_text(sid)
        sections.append(n)

    return {"book": book, "chapter": chapter, "sections": sections,
            "flat_text": ""}


# -- DOC (legacy binary, Latin-1 fallback) -------------------------------------


def extract_doc(filepath: str) -> str:
    """Extract readable text from a binary .doc (OLE Word) file.

    Lossy Latin-1 decode -- no olefile dependency. Returns raw text.
    """
    try:
        with open(filepath, "rb") as f:
            raw = f.read()
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("cannot read DOC: %s" % exc) from exc

    text = raw.decode("latin-1").replace("\x00", "")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    runs = re.findall(r"[\x20-\x7e\t\r\n]{20,}", text)

    clean_lines = []
    for run in runs:
        run = run.strip()
        if not run:
            continue
        alpha = sum(1 for ch in run if ch.isalpha())
        if alpha / len(run) < 0.3:
            continue
        clean_lines.append(re.sub(r"[ \t]+", " ", run))

    return ("(extracted via binary Latin-1 fallback -- formatting lost)\n\n"
            + "\n".join(clean_lines))
